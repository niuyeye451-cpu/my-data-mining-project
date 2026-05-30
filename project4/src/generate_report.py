"""
Generate Experiment Report (.docx) — Power Load Forecasting
=============================================================
Rich report with extensive code snippets, embedded tables and figures.
"""

import os
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.join(os.path.dirname(__file__), "..")
OUT  = os.path.join(BASE, "output")

# ===================== Helpers =====================
def sfont(run, name="宋体", size=Pt(12), bold=False, color=None):
    run.font.size = size
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def scell(cell, text, name="宋体", size=Pt(10), bold=False, center=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    sfont(run, name, size, bold)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

def body(doc, text, indent=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    sfont(r, "宋体", Pt(12))
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(4)

def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {qn("w:fill"): "F0F0F0", qn("w:val"): "clear"})
    pPr.append(shd)
    for i, line in enumerate(code_text.strip().split("\n")):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)

def add_image(doc, path, caption, width=Inches(5.2)):
    """Insert a centered image with caption. Skips if file missing."""
    if not os.path.exists(path):
        body(doc, f"[图片缺失: {os.path.basename(path)}]")
        return
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=width)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(caption); sfont(r2, "宋体", Pt(10))
    doc.add_paragraph()  # spacer

# ===================== Load Results =====================
metrics_df = pd.read_csv(os.path.join(OUT, "final_comparison.csv"), index_col=0)

# ===================== Build Report =====================
def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ============ COVER ============
    for text, size in [("重 庆 大 学", 22), ("学 生 实 验 报 告", 18)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); sfont(r, "黑体", Pt(size), bold=True)

    doc.add_paragraph()
    info = [
        ("实验课程名称", "数据挖掘"),
        ("实验名称", "空调用电负荷预测"),
        ("学    院", "大数据与软件学院"),
    ]
    tbl = doc.add_table(rows=len(info), cols=3, style="Table Grid")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lab, val) in enumerate(info):
        scell(tbl.cell(i, 0), lab, "宋体", Pt(12))
        tbl.cell(i, 1).merge(tbl.cell(i, 2))
        scell(tbl.cell(i, 1), val, "宋体", Pt(12))
        tbl.cell(i, 0).width = Cm(3)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("实验  空调用电负荷预测"); sfont(r, "黑体", Pt(16), bold=True)

    # ============ 一、实验目的 ============
    doc.add_page_break()
    heading(doc, "一、实验目的")
    for t in [
        "1. 理解时间序列预测中数据预处理、特征工程的关键作用；",
        "2. 掌握随机森林（RF）、XGBoost、LSTM的算法原理及其在电力负荷预测中的应用；",
        "3. 实现MLP-RF融合模型，对比分析不同算法的预测性能；",
        "4. 通过MSE、MAE、RMSE、MAPE、R²五种指标综合评估模型表现；",
        "5. 分析不同模型在时间序列任务上的适用性和局限性。",
    ]:
        body(doc, t, indent=False)

    # ============ 二、实验原理 ============
    heading(doc, "二、实验原理")

    heading(doc, "2.1 数据预处理与特征工程", level=2)
    body(doc, "本实验使用UCI Individual Household Electric Power Consumption数据集，该数据集记录了法国巴黎某家庭2006年12月至2010年11月共47个月的分钟级能耗数据，包含7个变量：Global_active_power（总有功功率）、Global_reactive_power（总无功功率）、Voltage（电压）、Global_intensity（电流）、Sub_metering_1（厨房）、Sub_metering_2（洗衣间）、Sub_metering_3（空调/热水器）。")
    body(doc, "预处理流程包括：（1）缺失值处理（1.25%行缺失，直接删除）；（2）逐小时均值重采样（2,075,259行→34,168行）；（3）Min-Max归一化至[0,1]区间；（4）时间特征构建（hour_sin/cos循环编码保留日周期的循环特性、month_sin/cos编码年周期、is_weekend区分工作日/周末）；（5）滞后特征（lag1h/lag2h/lag3h/lag6h/lag12h/lag24h共6个）和滚动统计特征（6h/12h/24h的均值、标准差、最小值、最大值，共12个）。最终得到30维特征向量。")

    heading(doc, "2.2 随机森林（Random Forest）", level=2)
    body(doc, "随机森林是一种基于Bagging思想的集成学习方法，通过构建多个决策树并对其预测结果取平均来提高精度并控制过拟合。其核心优势在于：（1）通过随机样本采样和随机特征选择增加树的多样性；（2）天然抗噪声、不易过拟合；（3）能处理高维数据且无需特征归一化。本实验使用scikit-learn的RandomForestRegressor，通过RandomizedSearchCV进行超参数调优（15轮×3折交叉验证），最终选定n_estimators=200、max_depth=25、min_samples_split=5、min_samples_leaf=4、max_features=0.5。")

    heading(doc, "2.3 XGBoost", level=2)
    body(doc, "XGBoost（eXtreme Gradient Boosting）是一种基于梯度提升框架的集成学习算法。与随机森林不同，XGBoost的树是串行构建的，每棵新树拟合前树的残差。其核心优势包括：（1）在损失函数中引入正则项（reg_alpha/reg_lambda）防止过拟合；（2）支持特征并行处理（tree_method='hist'）和缺失值自动学习；（3）采用反向剪枝策略避免局部最优。本实验使用xgboost库的XGBRegressor，经RandomizedSearchCV调优（20轮×3折），最终选定n_estimators=400、max_depth=6、learning_rate=0.05、subsample=0.8、colsample_bytree=0.7。")

    heading(doc, "2.4 LSTM（长短期记忆网络）", level=2)
    body(doc, "LSTM是循环神经网络（RNN）的变体，通过引入遗忘门、输入门和输出门三个门控机制，有效解决了传统RNN在处理长序列时的梯度消失和梯度爆炸问题。本实验使用PyTorch构建3层LSTM网络（hidden_dim=128, dropout=0.3），以过去24小时的电力负荷序列（10维多变量特征）预测下一小时负荷。训练时使用Adam优化器（lr=1e-3）、ReduceLROnPlateau学习率衰减、梯度裁剪（max_norm=1.0）和早停策略（patience=15）。")

    heading(doc, "2.5 MLP-RF融合模型", level=2)
    body(doc, "MLP-RF融合模型将多层感知机（MLP）与随机森林（RF）的预测结果进行集成。MLP是一个3层全连接神经网络（256→128→64），使用ReLU激活、BatchNorm归一化和Dropout（0.3）正则化。融合策略采用简单平均法——对两个基模型的输出取算术平均。实验证明，简单平均融合（R²=0.9007）优于基于Ridge回归的Stacking融合（R²=0.8905），因为两个模型的预测值高度相关，线性元学习器无法从中提取额外信息，而简单平均使得两个模型的不同误差模式部分抵消。")

    # ============ 三、关键代码及注释 ============
    doc.add_page_break()
    heading(doc, "三、关键代码及注释")

    # --- 3.1 ---
    heading(doc, "3.1 数据加载与重采样", level=2)
    body(doc, "使用内存优化策略加载127MB的原始数据：指定列加载、float32数据类型、立即逐小时重采样（2M→35K行），将内存占用从~900MB降至~2MB。")
    add_code(doc, """# 仅加载需要的列，指定float32减少内存
USECOLS = ["Date", "Time", "Global_active_power",
           "Global_reactive_power", "Voltage", "Global_intensity",
           "Sub_metering_1", "Sub_metering_2", "Sub_metering_3"]
DTYPES = {col: "float32" for col in USECOLS if col not in ("Date", "Time")}

df = pd.read_csv("household_power_consumption.txt",
                 sep=";", usecols=USECOLS, dtype=DTYPES,
                 low_memory=False, na_values="?")

# 合并日期时间并设为索引
df["datetime"] = pd.to_datetime(df["Date"] + " " + df["Time"], dayfirst=True)
df.drop(columns=["Date", "Time"], inplace=True)
df.set_index("datetime", inplace=True)

# 逐小时均值重采样 —— 大幅减少数据量
df = df.resample("1h").mean()  # 2,075,259 → 34,168 行
df.dropna(inplace=True)""")

    # --- 3.2 ---
    heading(doc, "3.2 时间特征与循环编码", level=2)
    body(doc, "构建时间特征并使用sin/cos循环编码保留周期性（例如23时与0时在圆上相邻，而非数值上的23→0跳变）：")
    add_code(doc, """h = df.index
df["hour"]       = h.hour.astype("int16")
df["dayofweek"]  = h.dayofweek.astype("int16")
df["month"]      = h.month.astype("int16")
df["is_weekend"] = (df["dayofweek"] >= 5).astype("int16")

# sin/cos循环编码：保留日周期（24h）和年周期（12月）的循环性质
df["hour_sin"]  = np.sin(2 * np.pi * df["hour"] / 24).astype("float32")
df["hour_cos"]  = np.cos(2 * np.pi * df["hour"] / 24).astype("float32")
df["month_sin"] = np.sin(2 * np.pi * df["month"] / 12).astype("float32")
df["month_cos"] = np.cos(2 * np.pi * df["month"] / 12).astype("float32")""")

    # --- 3.3 ---
    heading(doc, "3.3 滞后特征与滚动统计", level=2)
    body(doc, "滞后特征直接编码时序自相关，滚动统计捕捉局部窗口的分布特征：")
    add_code(doc, """# 滞后特征：过去时刻的负荷值
for lag in [1, 2, 3, 6, 12, 24]:
    df[f"lag{lag}h"] = df[TARGET].shift(lag)

# 滚动窗口统计：6h/12h/24h 的均值、标准差、最小值、最大值
for w in [6, 12, 24]:
    shifted = df[TARGET].shift(1)
    df[f"roll_mean_{w}h"] = shifted.rolling(w).mean()
    df[f"roll_std_{w}h"]  = shifted.rolling(w).std()
    df[f"roll_min_{w}h"]  = shifted.rolling(w).min()
    df[f"roll_max_{w}h"]  = shifted.rolling(w).max()

# 去除因滞后/滚动产生的NaN行
df.dropna(inplace=True)  # 最终 34,144 行""")

    # --- 3.4 ---
    heading(doc, "3.4 数据泄漏防护", level=2)
    body(doc, "在初步实验中，发现Global_intensity（电流）特征的重要性高达99.9%。分析发现这是因为物理公式P=V×I（功率=电压×电流），电压几乎恒定在240V，因此电流与目标变量Global_active_power几乎完全线性相关。这是典型的数据泄漏——模型不是在「预测」，而是在「读取答案」。")
    add_code(doc, """# 移除与目标变量存在物理公式关联的泄漏特征
EXCLUDE = {"Global_active_power",      # 目标变量本身
           "Global_intensity",          # P = V×I, V≈constant → I∝P
           "Global_reactive_power"}     # 无功功率也与有功功率强相关
feature_cols = [c for c in all_columns if c not in EXCLUDE]""")

    # --- 3.5 ---
    heading(doc, "3.5 随机森林训练与调参", level=2)
    body(doc, "使用RandomizedSearchCV进行高效超参数搜索，3折交叉验证防止过拟合：")
    add_code(doc, """from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import RandomizedSearchCV

rf_base = RandomForestRegressor(random_state=42, n_jobs=1)  # 单线程,CV层并行
rf_grid = {
    "n_estimators":      [100, 200, 300],
    "max_depth":         [10, 15, 20, 25, None],
    "min_samples_split": [2, 5, 10],
    "min_samples_leaf":  [1, 2, 4],
    "max_features":      [0.5, 0.7, 1.0],
}
rf_search = RandomizedSearchCV(
    rf_base, rf_grid, n_iter=15, cv=3,
    scoring="neg_mean_absolute_error", random_state=42, n_jobs=-1
)
rf_search.fit(X_train, y_train)

# 在训练+验证集上用最优参数重新训练
rf_final = RandomForestRegressor(**rf_search.best_params_,
                                 random_state=42, n_jobs=-1)
rf_final.fit(np.concatenate([X_train, X_val]),
             np.concatenate([y_train, y_val]))""")

    # --- 3.6 ---
    heading(doc, "3.6 XGBoost训练与调参", level=2)
    add_code(doc, """import xgboost as xgb

xgb_base = xgb.XGBRegressor(
    objective="reg:squarederror", random_state=42,
    n_jobs=-1, tree_method="hist"  # hist方法大幅加速
)
xgb_grid = {
    "n_estimators":      [100, 200, 400, 600],
    "max_depth":         [4, 6, 8, 10, 12],
    "learning_rate":     [0.01, 0.03, 0.05, 0.1],
    "subsample":         [0.7, 0.8, 0.9, 1.0],
    "colsample_bytree":  [0.7, 0.8, 0.9, 1.0],
    "reg_alpha":         [0, 0.1, 1.0],     # L1正则化
    "reg_lambda":        [1.0, 1.5, 2.0],   # L2正则化
}
xgb_search = RandomizedSearchCV(
    xgb_base, xgb_grid, n_iter=20, cv=3,
    scoring="neg_mean_absolute_error", random_state=42, n_jobs=-1
)
xgb_search.fit(X_train, y_train)""")

    # --- 3.7 ---
    heading(doc, "3.7 LSTM模型定义", level=2)
    body(doc, "使用PyTorch构建3层LSTM，输入为过去24小时的10维多变量序列，输出为下一小时的负荷预测值。包含梯度裁剪、学习率衰减和早停机制。")
    add_code(doc, """import torch, torch.nn as nn

class LSTMForecaster(nn.Module):
    def __init__(self, input_dim, hidden_dim=128, num_layers=3, dropout=0.3):
        super().__init__()
        self.lstm = nn.LSTM(input_dim, hidden_dim, num_layers,
                            batch_first=True, dropout=dropout)
        self.regressor = nn.Sequential(
            nn.Linear(hidden_dim, 64), nn.ReLU(),
            nn.Dropout(dropout), nn.Linear(64, 1),
        )

    def forward(self, x):               # x: (batch, 24, 10)
        out, _ = self.lstm(x)           # → (batch, 24, 128)
        out = out[:, -1, :]             # 取最后一个时间步
        return self.regressor(out).squeeze(-1)

# 训练循环（精简）
for epoch in range(EPOCHS):
    model.train()
    for xb, yb in train_loader:
        xb, yb = xb.to(DEVICE), yb.to(DEVICE)
        optimizer.zero_grad()
        loss = criterion(model(xb), yb)
        loss.backward()
        torch.nn.utils.clip_grad_norm_(model.parameters(), 1.0)  # 梯度裁剪
        optimizer.step()

    val_loss = evaluate(model, val_loader)
    scheduler.step(val_loss)            # 学习率衰减
    if val_loss < best_loss:
        best_loss = val_loss; torch.save(model.state_dict(), "lstm_best.pt")
    else:
        patience_left -= 1
        if patience_left == 0: break     # 早停""")

    # --- 3.8 ---
    heading(doc, "3.8 MLP-RF融合模型", level=2)
    body(doc, "MLP使用3层全连接网络，融合采用简单平均策略（实验证明优于Stacking）：")
    add_code(doc, """# MLP 定义
class MLP(nn.Module):
    def __init__(self, in_dim, hidden=[256, 128, 64], dropout=0.3):
        super().__init__()
        layers = []
        prev = in_dim
        for h in hidden:
            layers.extend([nn.Linear(prev, h), nn.BatchNorm1d(h),
                          nn.ReLU(), nn.Dropout(dropout)])
            prev = h
        layers.append(nn.Linear(prev, 1))
        self.net = nn.Sequential(*layers)

    def forward(self, x):
        return self.net(x).squeeze(-1)

# 简单平均融合 —— 优于Ridge Stacking
rf_pred = rf_model.predict(X_test)
mlp_pred = mlp_model(X_test_tensor).cpu().numpy()
ensemble_pred = (rf_pred + mlp_pred) / 2.0  # 简单算术平均""")

    # --- 3.9 ---
    heading(doc, "3.9 评估指标计算", level=2)
    body(doc, "统一使用五个回归指标进行全面评估：")
    add_code(doc, """def metrics(y_true, y_pred):
    y_t, y_p = y_true.ravel(), y_pred.ravel()
    mse  = np.mean((y_t - y_p) ** 2)                            # 均方误差
    mae  = np.mean(np.abs(y_t - y_p))                           # 平均绝对误差
    rmse = np.sqrt(mse)                                         # 均方根误差
    mape = np.mean(np.abs((y_t - y_p) / np.clip(y_t,1e-8,None))) * 100
    r2   = 1 - np.sum((y_t - y_p)**2) / np.sum((y_t - y_t.mean())**2)
    return {"MSE": mse, "MAE": mae, "RMSE": rmse,
            "MAPE(%)": mape, "R²": r2}""")

    # ============ 四、实验结果 ============
    doc.add_page_break()
    heading(doc, "四、实验结果")

    heading(doc, "4.1 实验设置", level=2)
    for s in [
        "数据集：UCI Individual Household Electric Power Consumption",
        "时间范围：2006-12-16 ~ 2010-11-26（47个月），逐小时重采样后34,168条",
        "数据划分：训练集 2006-12~2009-09（70%, 23,900条），验证集 2009-09~2010-04（15%, 5,122条），测试集 2010-04~2010-11（15%, 5,122条）",
        "特征维度：30维（7原始 + 8时间 + 6滞后 + 12滚动统计 − 2泄漏特征 + 3子计量）",
        "硬件环境：NVIDIA RTX 4060 Laptop (8GB) + PyTorch 2.12 + CUDA 13.0",
        "评估指标：MSE / MAE / RMSE / MAPE / R²",
    ]:
        body(doc, "  " + s, indent=False)

    # ---- 4.2 模型性能对比表 ----
    heading(doc, "4.2 模型性能对比", level=2)
    body(doc, "表1  五模型性能对比（测试集，时序对齐，按R²降序排列）", indent=False)
    doc.add_paragraph()

    headers = ["模型", "MSE", "MAE", "RMSE", "MAPE(%)", "R²"]
    n_rows = len(metrics_df) + 1
    tbl = doc.add_table(rows=n_rows, cols=6, style="Table Grid")
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        scell(tbl.cell(0, j), h, "宋体", Pt(9), bold=True)

    for i, (idx, row) in enumerate(metrics_df.iterrows()):
        scell(tbl.cell(i+1, 0), idx, "宋体", Pt(9))
        scell(tbl.cell(i+1, 1), f'{row["MSE"]:.4f}', "宋体", Pt(9))
        scell(tbl.cell(i+1, 2), f'{row["MAE"]:.4f}', "宋体", Pt(9))
        scell(tbl.cell(i+1, 3), f'{row["RMSE"]:.4f}', "宋体", Pt(9))
        scell(tbl.cell(i+1, 4), f'{row["MAPE(%)"]:.2f}', "宋体", Pt(9))
        scell(tbl.cell(i+1, 5), f'{row["R²"]:.4f}', "宋体", Pt(9))

    doc.add_paragraph()
    body(doc, f"最佳R²模型：{metrics_df.index[0]}（R²={metrics_df.iloc[0]['R²']:.4f}）；最佳MAPE模型：{metrics_df['MAPE(%)'].idxmin()}（MAPE={metrics_df['MAPE(%)'].min():.2f}%）。", indent=False)

    # ---- 4.2b 综合对比图 ----
    add_image(doc, os.path.join(OUT, "final_summary.png"),
              "图1  五模型综合对比（MAPE/R²/RMSE柱状图 + 预测曲线 + 关键发现）")

    # ---- 4.3 EDA ----
    heading(doc, "4.3 探索性数据分析（EDA）", level=2)

    body(doc, "4.3.1 变量相关性分析", indent=False)
    body(doc, "计算7个原始变量之间的Pearson相关系数矩阵。Global_active_power与Global_intensity的相关系数接近1.0（物理公式P=V×I决定），验证了排除Global_intensity的合理性。Sub_metering_3（空调/热水器）与目标变量的相关性最高（~0.7），是预测的核心变量。")
    add_image(doc, os.path.join(OUT, "eda_correlation_heatmap.png"),
              "图2  变量相关性矩阵热力图", Inches(4.5))

    body(doc, "4.3.2 季节性分解", indent=False)
    body(doc, "对逐日均值序列进行Seasonal Decomposition（周期=7天），分解为趋势（Trend）、季节性（Seasonal）和残差（Residual）。可以观察到：（1）长期趋势在2008-2009年有上升后回落；（2）周季节性模式明显（周末用电高于工作日）；（3）残差基本呈现白噪声特征。")
    add_image(doc, os.path.join(OUT, "eda_seasonal_decompose.png"),
              "图3  季节性分解（Trend + Seasonal + Residual, 周期=7天）", Inches(5.0))

    body(doc, "4.3.3 功率变化热力图", indent=False)
    body(doc, "以日期为行、小时为列绘制Global_active_power的热力图（最后60天）。颜色越深表示功率越高。可以清晰看到：（1）白天（10:00-18:00）功率显著高于夜间；（2）周末的用电模式与工作日不同；（3）日间存在两个峰值（午间和傍晚），对应做饭和回家后的用电高峰。")
    add_image(doc, os.path.join(OUT, "eda_power_heatmap.png"),
              "图4  功率变化热力图（日期×小时，最后60天）", Inches(5.0))

    # ---- 4.4 预测结果可视化 ----
    heading(doc, "4.4 预测结果可视化", level=2)

    body(doc, "4.4.1 五模型预测曲线对比（前7天）", indent=False)
    body(doc, "选取测试集前7天（168小时），同时展示观测值和五个模型的预测值。MLP-RF Avg和XGBoost最贴合真实曲线，LSTM的预测偏差明显偏大，尤其在负荷突变处（如早晚高峰）响应迟缓。")
    add_image(doc, os.path.join(OUT, "all_models_7days.png"),
              "图5  五模型预测曲线对比（测试集前7天）", Inches(5.2))

    body(doc, "4.4.2 预测值vs真实值散点图", indent=False)
    body(doc, "每个点代表一个测试样本，横轴为真实值，纵轴为预测值，越接近对角线y=x表示预测越准确。树模型（RF/XGBoost）和MLP的点紧密围绕对角线，而LSTM在低值区（夜间低负荷）和高值区（日间高峰）均存在较大偏差。")
    add_image(doc, os.path.join(OUT, "all_models_scatter.png"),
              "图6  五模型预测值vs真实值散点图", Inches(5.2))

    body(doc, "4.4.3 模型残差分布", indent=False)
    body(doc, "RF和XGBoost的残差（真实值−预测值）分布均呈近似正态、以零为中心、标准差较小（~0.036），表明模型不存在系统性偏差。XGBoost的残差分布略窄于RF，与MAPE指标一致。")
    add_image(doc, os.path.join(OUT, "rf_xgb_residuals.png"),
              "图7  RF与XGBoost残差分布图", Inches(5.0))

    body(doc, "4.4.4 LSTM训练曲线", indent=False)
    body(doc, "LSTM在约第48轮达到最优（早停于第63轮），验证集MSE最低约0.0067，但远高于RF/XGBoost的~0.0013，与最终测试性能（R²=0.579）一致。训练损失持续下降但验证损失在~0.007附近停滞，表明模型已达到容量上限。")
    add_image(doc, os.path.join(OUT, "lstm_training_curves.png"),
              "图8  LSTM训练与验证损失曲线", Inches(4.5))

    body(doc, "4.4.5 三模型全测试期对比", indent=False)
    body(doc, "全测试期（2010年4月至11月，约7个月）的预测曲线。三种模型均能较好地跟踪负荷的长期趋势和季节变化，MLP-RF Avg的融合预测（红色虚线）整体最贴合观测值，尤其在负荷突变区域保持了更好的跟踪能力。")
    add_image(doc, os.path.join(OUT, "all_models_barchart.png"),
              "图9  五模型性能柱状图对比（MAPE & R²）", Inches(5.0))

    # ---- 4.5 特征重要性 ----
    heading(doc, "4.5 特征重要性分析", level=2)
    body(doc, "随机森林和XGBoost的特征重要性排序高度一致：")
    body(doc, "• Sub_metering_3（空调/热水器能耗）：重要性~38%，是预测电力负荷的最强特征——空调是家庭负荷的主要驱动力。")
    body(doc, "• lag1h（前一小时负荷）：重要性~12-21%，反映了电力负荷的强时序自相关性——最近时刻的负荷是预测下一时刻最直接的信息。")
    body(doc, "• Sub_metering_1 & Sub_metering_2（厨房和洗衣间）：合计贡献约20%，特定电器的使用模式提供了额外的预测信号。")
    body(doc, "• Voltage（电压）：重要性~2-3%，电网电压的微小波动也包含一定的预测信息。")
    body(doc, "• 时间循环特征（hour_sin/cos, month_cos）：日周期和年周期特征合计贡献约5%，为模型提供了时间上下文。")

    # ============ 五、实验分析 ============
    doc.add_page_break()
    heading(doc, "五、实验分析")

    heading(doc, "5.1 模型性能排序与分析", level=2)
    body(doc, "五模型在测试集上的R²排序为：MLP-RF Avg (0.9007) > XGBoost (0.8947) > Random Forest (0.8937) > MLP (0.8904) > LSTM (0.5793)。")
    body(doc, "（1）MLP-RF简单平均融合取得最优R²，证明了模型集成的有效性。RF和MLP虽是不同算法范式（树模型vs神经网络），但错误模式不完全重合，简单平均使得两者的偏差部分抵消，R²从RF的0.8937提升至0.9007。")
    body(doc, "（2）XGBoost在MAPE指标上最优（23.79%），在实际部署场景中更具价值——MAPE直观反映预测误差百分比，是业务人员最关注的指标。XGBoost的训练效率（106s）也远优于RF（1201s），工程实用性更强。")
    body(doc, "（3）树模型（RF/XGBoost）表现稳健，主要得益于丰富的显式特征工程（滞后+滚动统计），这些特征直接编码了时序依赖关系，降低了模型从原始数据中学习的难度。")
    body(doc, "（4）MLP（3层，256→128→64，BatchNorm+Dropout）表现与树模型接近（R²=0.8904），且训练时间仅23秒，证明了深度神经网络在表格化时序数据上同样具有竞争力。BatchNorm加速了训练收敛，Dropout有效防止了过拟合。")

    heading(doc, "5.2 LSTM性能不佳的原因分析", level=2)
    body(doc, "LSTM在本次实验中表现最差（MAPE=53.9%, R²=0.579），主要原因有三：")
    body(doc, "（1）特征不对称：LSTM仅使用10维原始特征序列（含自回归项），而RF/XGBoost获得30维显式工程特征（滞后+滚动统计）。LSTM需要从序列中隐式学习这些模式，对模型容量和数据量的要求更高。")
    body(doc, "（2）数据量不足：~29K训练样本对3层LSTM（~200K参数）而言偏少。相比之下，RF/XGBoost基于树的集成方法在小样本场景下天然具有优势。")
    body(doc, "（3）任务特性：对于已做好特征工程的表格化时序预测，梯度提升树通常是更优选择。LSTM的优势在于处理原始长序列、强非线性多变量交互，或在大数据量下端到端学习。")

    heading(doc, "5.3 数据泄漏的发现与修复", level=2)
    body(doc, "实验初期，模型R²曾达到0.999，表面上看近乎完美。但特征重要性分析揭示Global_intensity（电流）占据了99.9%的重要性——这是典型的数据泄漏。根据物理定律P=V×I，在电压恒定（~240V）的情况下，电流与功率几乎成完美线性关系。将电流作为特征预测功率，相当于直接给出了答案。")
    body(doc, "这一教训提醒我们：在特征工程中，必须深入理解数据的物理含义和变量间的因果关系，不能盲目追求高指标。R²=0.999不是庆祝的理由，而是检查数据泄漏的警报。一个简单的检查方法：如果某个特征的importance超过90%，应当立即审视其与目标变量的物理/逻辑关系。")

    heading(doc, "5.4 Stacking与简单平均的对比", level=2)
    body(doc, "实验对比了两种融合策略：基于Ridge回归的Stacking和简单算术平均。结果发现Stacking的Ridge元学习器给MLP的权重为-0.07（基本为零），融合效果退化为单独使用RF（R²=0.8905）。而简单平均却取得了最优R²=0.9007。")
    body(doc, "原因分析：当两个基模型的预测高度相关时（RF和MLP在测试集上的预测相关系数>0.95），线性元学习器无法从两者的线性组合中获得新信息——Ridge回归几乎完全依赖其中一个模型，失去了集成的意义。相反，简单平均强制给两个模型各50%的权重，恰好使得两个模型的不同偏差模式部分抵消，产生了更好的泛化效果。这启示我们：在基模型预测高度相关时，简单平均可能比复杂的元学习更有效。")

    # ============ 六、总结 ============
    heading(doc, "六、实验总结")
    body(doc, "本次实验系统地完成了空调用电负荷预测的完整数据挖掘流程——从原始数据加载、探索性分析、特征工程，到五种模型的训练、超参数调优与综合评估对比。主要结论如下：")
    body(doc, "（1）XGBoost是实际部署的最佳选择，在预测精度（MAPE=23.8%）和训练效率（106s）之间取得了最优平衡。")
    body(doc, "（2）MLP-RF简单平均融合取得了最高R²（0.9007），证明了模型集成在时序预测中的价值，但也揭示了当基模型预测高度相关时，Stacking的局限性。")
    body(doc, "（3）深度LSTM在本任务中表现不佳（MAPE=53.9%, R²=0.579），主要受限于特征维度和数据量。这提醒我们：深度学习并非万能，任务的特性（数据规模、特征结构）决定了最合适的算法。")
    body(doc, "（4）数据泄漏的发现和修复是本次实验最有价值的教训之一：物理理解比盲目追求高指标更重要。一个R²=0.999的「完美模型」可能只是学会了答案。")
    body(doc, "（5）特征工程是决定模型性能上限的关键——Sub_metering_3（空调分项计量，重要性~38%）和lag1h（时序自相关，重要性~21%）是预测力最强的两个特征。")
    body(doc, "（6）对于表格化的时序预测任务，经过充分特征工程后，梯度提升树方法（XGBoost）仍然是工业界最可靠的首选方案。")
    body(doc, "")
    body(doc, "通过本次实验，对时间序列预测的完整方法论、多种机器学习算法的特性与适用场景、以及模型评估与融合策略有了深入的理解和实践经验。")

    # ============ Save ============
    out_path = os.path.join(OUT, "空调负荷预测实验报告.docx")
    doc.save(out_path)
    print(f"Report saved to: {out_path}")

if __name__ == "__main__":
    main()
